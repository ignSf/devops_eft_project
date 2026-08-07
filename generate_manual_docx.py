import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="CBD5E1"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_formatted_text(paragraph, text, default_font="Calibri", default_size=11, default_color=RGBColor(51, 65, 85)):
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        r = paragraph.add_run()
        r.font.name = default_font
        r.font.size = Pt(default_size)
        r.font.color.rgb = default_color
        
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            r.text = token[2:-2]
            r.bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) >= 2:
            r.text = token[1:-1]
            r.font.name = "Consolas"
            r.font.size = Pt(default_size - 0.5)
            r.font.color.rgb = RGBColor(14, 116, 144) # Teal
        else:
            r.text = token

def process_table_lines(doc, table_lines):
    if not table_lines:
        return
    
    # Parse rows
    parsed_rows = []
    for line in table_lines:
        # line starts and ends with '|'
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # Skip header separator line e.g. | --- | --- |
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        parsed_rows.append(cells)
    
    if not parsed_rows:
        return
    
    num_cols = max(len(row) for row in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "CBD5E1")
    
    for row_idx, row_data in enumerate(parsed_rows):
        is_header = (row_idx == 0)
        row = table.rows[row_idx]
        
        for col_idx in range(num_cols):
            cell = row.cells[col_idx]
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
            if is_header:
                set_cell_background(cell, "0F172A") # Dark Navy Header
            else:
                if row_idx % 2 == 1:
                    set_cell_background(cell, "FFFFFF")
                else:
                    set_cell_background(cell, "F8FAFC") # Soft zebra row
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            # Handle <br> in table cell text
            lines_in_cell = cell_text.split("<br>")
            for i, line_item in enumerate(lines_in_cell):
                if i > 0:
                    p = cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                
                font_color = RGBColor(255, 255, 255) if is_header else RGBColor(51, 65, 85)
                add_formatted_text(p, line_item, default_font="Calibri", default_size=10, default_color=font_color)
                if is_header:
                    for run in p.runs:
                        run.bold = True

def add_code_block(doc, code_text, lang=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "0F172A") # Slate dark background for code
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    lines = code_text.split("\n")
    for i, l in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        
        r = p.add_run(l)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(241, 245, 249) # Light off-white code text
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_callout(doc, quote_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "F0F9FF") # Soft blue callout
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Add left border accent
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:bottom w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_formatted_text(p, quote_text, default_font="Calibri", default_size=10.5, default_color=RGBColor(3, 105, 161))
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def generate_docx_for_file(md_file, output_filename):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_buffer = []
    code_lang = ""
    
    in_table = False
    table_buffer = []
    
    for line in lines:
        raw_line = line.rstrip("\r\n")
        stripped = raw_line.strip()
        
        # Handle Code Block Toggles
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                add_code_block(doc, "\n".join(code_buffer), code_lang)
                code_buffer = []
                in_code_block = False
            else:
                # Flush table if any
                if in_table:
                    process_table_lines(doc, table_buffer)
                    table_buffer = []
                    in_table = False
                in_code_block = True
                code_lang = stripped[3:].strip()
            continue
            
        if in_code_block:
            code_buffer.append(raw_line)
            continue
            
        # Handle Table Lines
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_buffer.append(stripped)
            continue
        else:
            if in_table:
                process_table_lines(doc, table_buffer)
                table_buffer = []
                in_table = False
                
        # Empty Line
        if not stripped:
            continue
            
        # Horizontal Rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("_________________________________________________________________________________")
            r.font.color.rgb = RGBColor(203, 213, 225)
            p.paragraph_format.space_after = Pt(6)
            continue
            
        # Blockquote / Callout
        if stripped.startswith("> "):
            quote_text = stripped[2:].strip()
            add_callout(doc, quote_text)
            continue
            
        # Headings
        if stripped.startswith("# "):
            h_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            add_formatted_text(p, h_text, default_font="Calibri", default_size=20, default_color=RGBColor(15, 23, 42))
            p.runs[0].bold = True
            
        elif stripped.startswith("## "):
            h_text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, h_text, default_font="Calibri", default_size=15, default_color=RGBColor(15, 23, 42))
            p.runs[0].bold = True
            
        elif stripped.startswith("### "):
            h_text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, h_text, default_font="Calibri", default_size=13, default_color=RGBColor(30, 41, 59))
            p.runs[0].bold = True
            
        elif stripped.startswith("#### "):
            h_text = stripped[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, h_text, default_font="Calibri", default_size=11.5, default_color=RGBColor(51, 65, 85))
            p.runs[0].bold = True
            
        # Images
        elif stripped.startswith("![") and "](" in stripped:
            match = re.search(r'!\[(.*?)\]\((.*?)\)', stripped)
            if match:
                alt_text, img_path = match.group(1), match.group(2)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.8))
                    
                    # Caption
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_after = Pt(10)
                    r_cap = cap_p.add_run(f"Figura: {alt_text}")
                    r_cap.font.name = "Calibri"
                    r_cap.font.size = Pt(9.5)
                    r_cap.font.italic = True
                    r_cap.font.color.rgb = RGBColor(100, 116, 139)
                else:
                    print(f"Warning: Image path not found: {img_path}")
                    
        # Bullet List Items
        elif stripped.startswith("* ") or stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            item_text = stripped[2:].strip()
            add_formatted_text(p, item_text, default_font="Calibri", default_size=10.5, default_color=RGBColor(51, 65, 85))
            
        # Numbered List Items
        elif re.match(r'^\d+[\.\)]\s+', stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            item_text = re.sub(r'^\d+[\.\)]\s+', '', stripped).strip()
            add_formatted_text(p, item_text, default_font="Calibri", default_size=10.5, default_color=RGBColor(51, 65, 85))
            
        # Normal Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, stripped, default_font="Calibri", default_size=11, default_color=RGBColor(51, 65, 85))

    # Flush any remaining table
    if in_table:
        process_table_lines(doc, table_buffer)
        
    doc.save(output_filename)
    print(f"Document saved successfully as: {output_filename}")

if __name__ == "__main__":
    generate_docx_for_file("MANUAL_EVIDENCIAS_PASO_A_PASO_AWS.md", "MANUAL_EVIDENCIAS_PASO_A_PASO_AWS.docx")
    generate_docx_for_file("GUIA_PASO_A_PASO_AWS_DEVOPS.md", "GUIA_PASO_A_PASO_AWS_DEVOPS.docx")

