import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Configuración de márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Título Principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("INFORME TÉCNICO DE EVALUACIÓN FINAL TRANSVERSAL")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    # Subtítulo
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Asignatura: ISY1101 - Introducción a Herramientas DevOps (Duoc UC 2025)")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(14, 116, 144)

    doc.add_paragraph() # Spacing

    # Leer Markdown
    with open("INFORME_TECNICO_DEVOPS.md", "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_list = False

    for line in lines:
        line_str = line.strip()
        
        if not line_str:
            continue

        if line_str.startswith("# "):
            continue # Título ya agregado

        elif line_str.startswith("## "):
            h_text = line_str.replace("## ", "")
            h = doc.add_heading(h_text, level=1)
            h.style.font.name = "Calibri"
            h.style.font.size = Pt(16)
            h.style.font.color.rgb = RGBColor(15, 23, 42)

        elif line_str.startswith("### "):
            h_text = line_str.replace("### ", "")
            h = doc.add_heading(h_text, level=2)
            h.style.font.name = "Calibri"
            h.style.font.size = Pt(13)
            h.style.font.color.rgb = RGBColor(30, 41, 59)

        elif line_str.startswith("* ") or line_str.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            text = line_str[2:]
            parts = text.split("**")
            for i, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True

        elif line_str[0].isdigit() and line_str[1:3] in [". ", ") "]:
            p = doc.add_paragraph(style="List Number")
            text = line_str[3:]
            parts = text.split("**")
            for i, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True

        elif line_str == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("_________________________________________________________________________________")
            r.font.color.rgb = RGBColor(203, 213, 225)

        elif line_str.startswith("![") and "](" in line_str:
            img_path = line_str.split("](")[1].rstrip(")")
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.8))

        else:
            p = doc.add_paragraph()
            parts = line_str.split("**")
            for i, part in enumerate(parts):
                r = p.add_run(part)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                if i % 2 == 1:
                    r.bold = True

    doc.save("INFORME_TECNICO_DEVOPS.docx")
    print("INFORME_TECNICO_DEVOPS.docx generado exitosamente.")

if __name__ == "__main__":
    create_report()
